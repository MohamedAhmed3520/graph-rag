"""Parse normalized loaded documents into text.

``load_documents`` in ``ingestion.loader`` reads raw bytes from disk. This
parser is responsible for turning those bytes into real, human-readable text
for each supported file type:

* PDF     -> text extracted page-by-page with ``pypdf``
* DOCX    -> paragraph/table text extracted with ``python-docx``
* TXT/MD  -> plain UTF-8 text

Previously PDFs were decoded as UTF-8 bytes, which produces the PDF's internal
binary structure (``%PDF-1.5``, xref tables, compressed streams) instead of the
actual lesson content. That garbage was then embedded, retrieved, and fed to
graph extraction, which is why retrieval reported "not enough info" and the
knowledge graph stayed empty even though the answer existed in the PDF.
"""
from __future__ import annotations

from ingestion.loader import LoadedDocument
from utils.logger import get_logger

logger = get_logger(__name__)


def _decode_text(content: bytes) -> str:
    """Decode raw bytes as text, falling back to a lossy decode.

    This is used for plain-text formats (``.txt``, ``.md``, ``.markdown``).
    """
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return content.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="ignore").strip()


def extract_pdf_text(content: bytes) -> str:
    """Extract text from a PDF byte stream using ``pypdf``."""
    from io import BytesIO

    from pypdf import PdfReader

    reader = PdfReader(BytesIO(content))

    pages: list[str] = []
    for page in reader.pages:
        try:
            page_text = page.extract_text() or ""
        except Exception as exc:  # pragma: no cover - defensive per-page handling
            logger.warning("Could not extract text from a PDF page: %s", exc)
            page_text = ""

        if page_text.strip():
            pages.append(page_text.strip())

    if not pages:
        logger.warning("PDF produced no extractable text.")

    return "\n".join(pages).strip()


def extract_docx_text(content: bytes) -> str:
    """Extract paragraph and table text from a DOCX byte stream."""
    from io import BytesIO

    from docx import Document

    document = Document(BytesIO(content))

    parts: list[str] = []

    # Paragraphs (body text, headings, etc.)
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text.strip())

    # Tables (documents often carry answers inside tables)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts).strip()


def parse_document(document: LoadedDocument) -> str:
    """Parse a single loaded document into clean text."""
    suffix = document.suffix

    try:
        if suffix == ".pdf":
            return extract_pdf_text(document.content)

        if suffix == ".docx":
            return extract_docx_text(document.content)

        if suffix in {".txt", ".md", ".markdown"}:
            return _decode_text(document.content)

        raise ValueError(f"Unsupported file type: {suffix}")

    except Exception as exc:
        logger.exception("Failed to parse %s", document.source_path)
        raise RuntimeError(f"Failed to parse {document.source_path}: {exc}") from exc
